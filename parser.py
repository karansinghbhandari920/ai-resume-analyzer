"""
parser.py
Two responsibilities:
  1. extract_text(...)  - pull raw text out of an uploaded PDF/DOCX/TXT file
  2. ResumeParser        - turn that raw text into structured fields (name,
                            email, phone, links, skills, education,
                            experience, projects, certifications)

spaCy is used for name detection (PERSON entities) and sentence
segmentation; everything else is deliberately regex/keyword based since
resumes are semi-structured documents and rules are more reliable than a
generic NER model for fields like email/phone/skills.
"""
from __future__ import annotations

import io
import re
from dataclasses import dataclass, field

import pdfplumber
import docx as docx_lib

from skills import SKILL_TAXONOMY

# --------------------------------------------------------------------------- #
# spaCy is optional at runtime: if the model hasn't been downloaded yet we
# fall back to regex-only name detection rather than crashing the app.
# --------------------------------------------------------------------------- #
try:
    import spacy

    try:
        _NLP = spacy.load("en_core_web_sm")
    except OSError:
        _NLP = None
except ImportError:
    _NLP = None


# --------------------------------------------------------------------------- #
# Text extraction
# --------------------------------------------------------------------------- #
def extract_text(uploaded_file) -> str:
    """uploaded_file is a Streamlit UploadedFile. Dispatches on extension."""
    name = uploaded_file.name.lower()
    raw_bytes = uploaded_file.read()
    uploaded_file.seek(0)

    if name.endswith(".pdf"):
        return _extract_pdf(raw_bytes)
    if name.endswith(".docx"):
        return _extract_docx(raw_bytes)
    if name.endswith(".txt"):
        return raw_bytes.decode("utf-8", errors="ignore")
    raise ValueError(f"Unsupported file type: {uploaded_file.name}. Use PDF, DOCX, or TXT.")


def _extract_pdf(raw_bytes: bytes) -> str:
    text_parts = []
    with pdfplumber.open(io.BytesIO(raw_bytes)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
    return "\n".join(text_parts)


def _extract_docx(raw_bytes: bytes) -> str:
    document = docx_lib.Document(io.BytesIO(raw_bytes))
    parts = [p.text for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(parts)


# --------------------------------------------------------------------------- #
# Regex patterns
# --------------------------------------------------------------------------- #
EMAIL_RE = re.compile(r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}")
PHONE_RE = re.compile(r"(\+?\d{1,3}[-.\s]?)?\(?\d{3,4}\)?[-.\s]?\d{3,4}[-.\s]?\d{3,4}")
LINKEDIN_RE = re.compile(r"(https?://)?(www\.)?linkedin\.com/in/[A-Za-z0-9_\-/]+", re.IGNORECASE)
GITHUB_RE = re.compile(r"(https?://)?(www\.)?github\.com/[A-Za-z0-9_\-/]+", re.IGNORECASE)

SECTION_HEADERS = {
    "summary": ["summary", "objective", "profile", "about me"],
    "education": ["education", "academic background", "qualifications"],
    "experience": ["experience", "work experience", "employment history", "professional experience"],
    "skills": ["skills", "technical skills", "core competencies"],
    "projects": ["projects", "personal projects", "academic projects"],
    "certifications": ["certifications", "certificates", "licenses"],
}

ACTION_VERBS = {
    "achieved", "administered", "analyzed", "automated", "built", "championed", "coordinated",
    "created", "delivered", "designed", "developed", "directed", "drove", "engineered",
    "established", "executed", "generated", "implemented", "improved", "increased", "initiated",
    "launched", "led", "managed", "optimized", "orchestrated", "organized", "pioneered",
    "produced", "reduced", "redesigned", "resolved", "spearheaded", "streamlined", "supervised",
    "transformed", "built", "deployed", "architected",
}


@dataclass
class ParsedResume:
    name: str | None = None
    email: str | None = None
    phone: str | None = None
    linkedin: str | None = None
    github: str | None = None
    skills: list[str] = field(default_factory=list)
    education: list[str] = field(default_factory=list)
    experience: list[str] = field(default_factory=list)
    projects: list[str] = field(default_factory=list)
    certifications: list[str] = field(default_factory=list)
    sections_found: list[str] = field(default_factory=list)
    bullet_points: list[str] = field(default_factory=list)

    def to_dict(self) -> dict:
        return {
            "name": self.name,
            "email": self.email,
            "phone": self.phone,
            "linkedin": self.linkedin,
            "github": self.github,
            "skills": self.skills,
            "education": self.education,
            "experience": self.experience,
            "projects": self.projects,
            "certifications": self.certifications,
            "sections_found": self.sections_found,
            "bullet_points": self.bullet_points,
        }


class ResumeParser:
    def __init__(self, text: str):
        self.text = text or ""
        self.lines = [l.strip() for l in self.text.splitlines() if l.strip()]

    def parse(self) -> ParsedResume:
        parsed = ParsedResume()
        parsed.email = self._find_email()
        parsed.phone = self._find_phone()
        parsed.linkedin = self._find_link(LINKEDIN_RE)
        parsed.github = self._find_link(GITHUB_RE)
        parsed.name = self._find_name()
        parsed.skills = self._find_skills()
        parsed.sections_found = self._find_sections()
        parsed.education = self._extract_section_lines("education")
        parsed.experience = self._extract_section_lines("experience")
        parsed.projects = self._extract_section_lines("projects")
        parsed.certifications = self._extract_section_lines("certifications")
        parsed.bullet_points = self._extract_bullets()
        return parsed

    # ---- individual extractors -------------------------------------- #
    def _find_email(self) -> str | None:
        match = EMAIL_RE.search(self.text)
        return match.group(0) if match else None

    def _find_phone(self) -> str | None:
        for match in PHONE_RE.finditer(self.text):
            digits = re.sub(r"\D", "", match.group(0))
            if 7 <= len(digits) <= 15:
                return match.group(0).strip()
        return None

    def _find_link(self, pattern: re.Pattern) -> str | None:
        match = pattern.search(self.text)
        return match.group(0) if match else None

    def _find_name(self) -> str | None:
        # Heuristic 1: spaCy PERSON entity near the top of the document
        if _NLP is not None:
            header = "\n".join(self.lines[:5])
            doc = _NLP(header)
            for ent in doc.ents:
                if ent.label_ == "PERSON" and 1 <= len(ent.text.split()) <= 4:
                    return ent.text.strip()
        # Heuristic 2: first non-empty line if it looks like "Firstname Lastname"
        for line in self.lines[:5]:
            words = line.split()
            if 1 < len(words) <= 4 and all(w.replace(".", "").isalpha() for w in words):
                if not any(c.isdigit() for c in line) and "@" not in line:
                    return line.title()
        return None

    def _find_skills(self) -> list[str]:
        found = []
        lowered = self.text.lower()
        for canonical, patterns in SKILL_TAXONOMY.items():
            for pat in patterns:
                if re.search(rf"(?<![A-Za-z0-9]){pat}(?![A-Za-z0-9])", lowered):
                    found.append(canonical)
                    break
        return found

    def _find_sections(self) -> list[str]:
        present = []
        lowered = self.text.lower()
        for section, aliases in SECTION_HEADERS.items():
            if any(re.search(rf"\b{re.escape(alias)}\b", lowered) for alias in aliases):
                present.append(section)
        return present

    def _extract_section_lines(self, section_key: str, max_lines: int = 12) -> list[str]:
        """Grabs lines that fall between this section's header and the next
        recognized header - a lightweight approach that works well on the
        single-column resumes most ATS-friendly templates use."""
        aliases = SECTION_HEADERS[section_key]
        all_aliases_flat = [a for aliases_list in SECTION_HEADERS.values() for a in aliases_list]

        start_idx = None
        for i, line in enumerate(self.lines):
            if any(re.fullmatch(rf"{re.escape(a)}s?", line.strip().lower()) for a in aliases) or \
               any(a in line.lower() and len(line) < 40 for a in aliases):
                start_idx = i + 1
                break
        if start_idx is None:
            return []

        collected = []
        for line in self.lines[start_idx:start_idx + 60]:
            if any(re.fullmatch(rf"{re.escape(a)}s?", line.strip().lower()) for a in all_aliases_flat):
                break
            collected.append(line)
            if len(collected) >= max_lines:
                break
        return collected

    def _extract_bullets(self) -> list[str]:
        bullets = []
        for line in self.lines:
            stripped = line.lstrip("•-*▪◦ \t")
            if stripped != line and len(stripped) > 3:
                bullets.append(stripped)
        # Fallback: if no bullet markers were used, treat reasonably long
        # sentences in the experience section as pseudo-bullets.
        if not bullets:
            bullets = [l for l in self.lines if 30 < len(l) < 220]
        return bullets

    def action_verb_ratio(self) -> float:
        bullets = self._extract_bullets()
        if not bullets:
            return 0.0
        hits = sum(1 for b in bullets if b.split()[0].lower().strip(".,") in ACTION_VERBS)
        return hits / len(bullets)

    def quantified_ratio(self) -> float:
        bullets = self._extract_bullets()
        if not bullets:
            return 0.0
        number_pattern = re.compile(r"\d")
        hits = sum(1 for b in bullets if number_pattern.search(b))
        return hits / len(bullets)
